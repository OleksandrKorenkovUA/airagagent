# Імпорт необхідних бібліотек для роботи з Milvus - векторною базою даних для пошуку схожих документів
from milvus_model.hybrid import BGEM3EmbeddingFunction
import re # для регулярних виразів
import subprocess
import numpy as np
from pymilvus import Collection

import torch
torch.classes.__path__ = []
import subprocess
from langchain_huggingface.llms import HuggingFacePipeline
from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline
# Імпорт бібліотек для роботи з LLM (великими мовними моделями)
from PIL import Image  # Для обробки зображень
from langchain_ollama import ChatOllama  # Для взаємодії з Ollama API
from langchain_core.prompts import ChatPromptTemplate  # Для створення шаблонів промптів
from langchain_text_splitters import RecursiveCharacterTextSplitter  # Для розбиття тексту на частини
# Імпорт для ранжування результатів пошуку
from pymilvus.model.reranker import BGERerankFunction  # Функція для переранжування результатів пошуку
# Імпорт для обробки тексту та зображень
from transformers import pipeline as hf_pipeline  # Конвеєр для обробки тексту та зображень
# Імпорт для роботи з документами різних форматів
import soundfile as sf  # Для читання PDF-файлів
from PyPDF2 import PdfReader  # Для читання PDF-файлів
from docx import Document as DocxDocument  # Для роботи з Word-документами
# Імпорт стандартних бібліотек Python
from datetime import datetime as dt  # Для роботи з датою та часом
from pathlib import Path  # Для роботи з шляхами файлів
from typing import Any, Dict  # Для типізації
import re  # Для роботи з регулярними виразами
import logging  # Для логування
import streamlit as st  # Для створення веб-інтерфейсу
import ollama  # Для взаємодії з Ollama API
import json  # Для роботи з JSON
import yt_dlp  # Для завантаження відео з YouTube
import hashlib  # Для створення хешів
import whisper  # Для транскрибації аудіо
import os  # Для роботи з операційною системою
import uuid  # Для генерації унікальних ідентифікаторів
import pandas as pd  # Для роботи з табличними даними
import base64  # Для кодування даних у base64
import io  # Для роботи з потоками введення-виведення
import sqlite3  # Для роботи з SQLite базами даних
from tqdm import tqdm  # Для відображення прогресу
from datetime import datetime  # Для роботи з датою та часом
from config import *  # Імпорт всіх змінних з файлу конфігурації
from rag import *  # Імпорт всіх функцій з файлу rag.py
# Налаштування логування для відстеження роботи програми
logger = logging.getLogger(__name__)

if 'chunk_size' not in st.session_state or st.session_state.chunk_size is None:
    st.session_state.chunk_size = CHUNK_SIZE
if 'chunk_overlap' not in st.session_state or st.session_state.chunk_overlap is None:
    st.session_state.chunk_overlap = OVERLAP
if 'ret_k_results' not in st.session_state or st.session_state.ret_k_results is None:
    st.session_state.ret_k_results = RET_K_RESULTS


# Ініціалізація моделі для створення ембедінгів (векторних представлень) тексту
EMBEDDER = BGEM3EmbeddingFunction(
    model_name=EMBEDING_MODEL,  # Назва моделі для ембедінгів
    device="cuda" if torch.cuda.is_available() else "cpu",        # Використання CPU для обчислень (найбезпечніший варіант)
    use_fp16=False       # Відключення використання половинної точності (fp16)
)

def reset_chat():
    """Скидає історію чату в сесії Streamlit, видаляючи всі повідомлення та стан привітання"""
    st.session_state.messages      = []  # Очищення списку повідомлень
    st.session_state.chat_history  = []  # Очищення історії чату
    st.session_state.pop("greeted", None)  # Видалення стану привітання


# функція для видалення дужок з тексту
def remove_think(text):
    return re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL) if '<think>' in text else text


def create_bge_m3_embeddings():
    """Створює функцію для генерації ембедінгів BGE-M3, повертаючи ініціалізований об'єкт EMBEDDER"""
    bge_m3_ef = EMBEDDER  # Використання глобального об'єкта EMBEDDER
    return bge_m3_ef

    
@st.cache_resource(show_spinner="Завантажуємо україномовну модель…")
def create_ukr_llm():
    model_id = UKRAINIAN_MODEL
    tokenizer = AutoTokenizer.from_pretrained(model_id)
    model = AutoModelForCausalLM.from_pretrained(
        model_id,
        device_map="auto",      
        torch_dtype=torch.bfloat16 if torch.cuda.is_available() else torch.float32
    )
    return pipeline(
        "text-generation",
        model=model,
        tokenizer=tokenizer,
        max_new_tokens=MAX_TOKENS,
        do_sample=False,
    )


def create_llm(model_name):
    """Створює екземпляр моделі LLM з вказаною назвою, налаштовуючи температуру для детермінованих відповідей"""
    llm = ChatOllama(
        model=model_name,  # Назва моделі для використання
        temperature=MODEL_TEMPERATURE,  # Температура 0 для детермінованих відповідей
        base_url=OLLAMA_URL,  # URL для взаємодії з Ollama API
    )
    return llm



def rerank_search(documents, query, limit=10):
    """Перераховує результати пошуку за допомогою BGE-reranker для покращення релевантності"""
    bge_rf = BGERerankFunction(
        model_name=RERANKER_MODEL,  # Модель для переранжування
        device="cuda" if torch.cuda.is_available() else "cpu",  # Використання CPU
        top_k=RERANK_K_RESULTS  # Кількість найкращих результатів для повернення
    )
    reranked_results = bge_rf(query, documents)  # Переранжування результатів
    return reranked_results

def create_prompt(system_prompt):
    """Створює шаблон промпту для чат-моделі з системним промптом та місцями для контексту і питання"""
    return ChatPromptTemplate.from_messages([
        ("system",
         system_prompt),  # Системний промпт для налаштування поведінки моделі
        ("human",
         "Context:\n{context}\n\nQuestion: {question}")  # Шаблон для запиту користувача
    ])


def create_stream_prompt(system_prompt, question, context):
    """Створює шаблон промпту для чат-моделі з системним промптом та місцями для контексту і питання"""
    prompt_text = (
    system_prompt +
    f"\n\nContext:\n{context}\n\nQuestion: {question}")
    return prompt_text



def create_chain(llm, prompt):
    """Створює ланцюжок обробки з LLM та промпту для послідовної обробки запитів"""
    chain = prompt | llm  # Об'єднання промпту та моделі в ланцюжок
    return chain

def get_llm_context(chain, context):
    """Отримує відповідь від моделі на основі наданого контексту без конкретного питання"""
    response = chain.invoke(
        {
            "text": context,  # Передача контексту в ланцюжок
        }
    )
    return response


def get_llm_response(chain, question, context):
    """Отримує відповідь від моделі на основі питання та контексту"""
    response = chain.invoke(
    {
        "question": question,  # Питання користувача
        "context": context,  # Контекст для відповіді
    })
    return response


def create_dir(dir_name):
    """Створює директорію з вказаною назвою, якщо вона не існує"""
    folder = Path(dir_name)  # Створення об'єкта Path
    folder.mkdir(parents=True, exist_ok=True)  # Створення директорії з батьківськими директоріями

def create_logger(logger_name = __name__):
    """Створює логер з вказаною назвою та рівнем логування INFO"""
    logger = logging.getLogger(logger_name)  # Створення логера
    logger.setLevel(logging.INFO)  # Встановлення рівня логування
    return logger

def extract_video_id(url: str):
    """Витягує ідентифікатор відео з URL YouTube за допомогою регулярного виразу"""
    youtube_regex = (
    r'(https?://)?(www\.)?'
    '(youtube|youtu|youtube-nocookie)\\.(com|be)/'
    '(watch\\?v=|embed/|v/|.+\\?v=)?([^&=%\\?]{11})')  # Регулярний вираз для пошуку ID відео
    match = re.match(youtube_regex, url)  # Пошук співпадіння
    if match:
        return match.group(6)  # Повернення ID відео
    return ""  # Повернення порожнього рядка, якщо ID не знайдено

def clean_filename(filename: str) -> str:
    """Очищення імені файлу від недопустимих символів для безпечного збереження"""
    # Видаляємо недопустимі символи для імен файлів
    return re.sub(r'[\\/*?:"<>|]', "", filename)  # Заміна недопустимих символів на порожній рядок

def query_ollama(prompt, model, image_b64=None):
    return ollama.chat(
        model=model,
        messages=[{"role":"user","content":prompt,"images": image_b64}],
        stream=True               # <- returns a generator
    )

def image_to_base64(image):
    """Конвертує зображення у формат base64 для передачі в API"""
    if image.mode == "RGBA":
        image = image.convert("RGB")  # Прибираємо альфа-канал для сумісності
    img_byte_arr = io.BytesIO()  # Створення буфера для зображення
    image.save(img_byte_arr, format="JPEG")  # Збереження зображення у форматі JPEG
    return base64.b64encode(img_byte_arr.getvalue()).decode("utf-8")  # Кодування в base64


def process_video(url: str, video_dir: str, model_size: str):
    """Обробляє відео з YouTube: завантажує, транскрибує та розбиває на частини для подальшого аналізу"""
    whisper_model = get_whisper(model_size)  # Отримання моделі Whisper (кешується)
    logger = logging.getLogger(__name__)  # Створення логера
    
    video_chunks = []  # Список для зберігання частин відео
    video_id = extract_video_id(url)  # Витягування ID відео
    unique_video_id = f"{video_id}_{hashlib.md5(url.encode()).hexdigest()[:8]}"  # Створення унікального ID
    if not video_id:
        raise ValueError("Неправильний URL відео YouTube")  # Перевірка коректності URL
    st.markdown(f"Обробка відео з URL: {url}")  # Логування початку обробки
    try:
        # Створюємо об'єкт Path для директорії
        video_dir_path = Path(video_dir)  # Створення об'єкта Path
        video_dir_path.mkdir(parents=True, exist_ok=True)  # Створення директорії

        with yt_dlp.YoutubeDL({'quiet': True}) as ydl:
            logger.info("Отримання інформації про відео...")  # Логування
            info = ydl.extract_info(url, download=False)  # Отримання інформації про відео
            # Отримуємо та очищуємо назву
            if st.session_state.video_name:
                title = st.session_state.video_name
            else:
                title = info.get("alt_title", "video")  # Отримання назви відео
            st.markdown(f"Назва відео: {title}")
            safe_title = clean_filename(title)  # Очищення назви
            unique_safe_title = f"{safe_title}_{hashlib.md5(url.encode()).hexdigest()[:8]}"  # Створення унікальної назви
            logger.info(f"Отримано назву відео: {title}, очищена назва: {safe_title}")  # Логування

            video_info = {
                    "title": title,  # Назва відео
                    "uploader": info.get("uploader", "Невідомий автор"),  # Автор відео
                    "duration": info.get("duration", 0),  # Тривалість відео
                    "upload_date": info.get("upload_date", ""),  # Дата завантаження
                    "view_count": info.get("view_count", 0),  # Кількість переглядів
                    "like_count": info.get("like_count", 0),  # Кількість лайків
                    "description": info.get("description", ""),  # Опис відео
                    "video_id": video_id}  # ID відео
            logger.info(f"Знайдено відео: {video_info['title']} ({video_info['duration']} сек)")  # Логування

            file_path = video_dir_path / f"{safe_title}.mp4"  # Шлях до відео
            final_audio_path = video_dir_path / f"{safe_title}.mp3"  # Шлях до аудіо
            transcript_path = video_dir_path / f"{safe_title}_transcript.txt"  # Шлях до транскрипції
            
            logger.info(f"Шляхи до файлів:")  # Логування
            logger.info(f"Відео: {file_path}")  # Логування
            logger.info(f"Аудіо: {final_audio_path}")  # Логування
            logger.info(f"Транскрипція: {transcript_path}")  # Логування

            ydl_opts = {
                'format': 'bestvideo[height<=720]+bestaudio/best[height<=720]',  # Формат відео
                'paths': {'home': str(video_dir_path)},  # Шлях для збереження
                'outtmpl': {'default': f'{safe_title}.%(ext)s'},  # Шаблон імені файлу
                'postprocessors': [
                    {'key': 'FFmpegVideoConvertor', 'preferedformat': 'mp4'},  # Конвертація у mp4
                    {'key': 'FFmpegExtractAudio', 'preferredcodec': 'mp3', 'preferredquality': '192'}  # Витягування аудіо
                ],
                'merge_output_format': 'mp4',  # Формат виходу
                'keepvideo': True,  # Зберігання відео
                'quiet': True,  # Тихий режим
                'noplaylist': True  # Без плейлистів
            }

            with yt_dlp.YoutubeDL(ydl_opts) as ydl_download:
                if not file_path.exists() or not final_audio_path.exists():
                    logger.info("Починаємо завантаження відео...")  # Логування
                    ydl_download.download([url])  # Завантаження відео
                    logger.info("Завантаження завершено")  # Логування
                else:
                    logger.info("Файли вже існують, пропускаємо завантаження")  # Логування

            if not file_path.exists():
                logger.error(f"Відео файл не знайдено після завантаження: {file_path}")  # Логування помилки
                raise FileNotFoundError(f"Відео файл не знайдено: {file_path}")  # Виклик помилки

            logger.info("Починаємо транскрибацію...")  # Логування
            result = whisper_model.transcribe(str(file_path))
            placeholder = st.empty()
              # Транскрибація відео
            transcript = result["text"]
            placeholder.text(transcript)   # Отримання тексту транскрипції
            logger.info(f"Транскрибація завершена, отримано {len(transcript)} символів")  # Логування
            
            with open(transcript_path, 'w', encoding='utf-8') as f:
                f.write(transcript)  # Запис транскрипції у файл
            logger.info(f"Транскрипцію збережено")  # Логування

            return  transcript, file_path, title, unique_video_id
    

    except Exception as e:
        logger.error(f"Помилка при обробці відео: {str(e)}", exc_info=True)  # Логування помилки
        raise


def build_video_json(
    txt: str,
    meta: Dict[str, Any],
    file_path: str,
    original_filename: str,
    chunk_size: int = st.session_state.chunk_size,
    overlap: int = st.session_state.chunk_overlap,
) -> Dict[str, Any]:
    """
    Перетворює розшифровку відео в структуру JSON для зберігання в Milvus,
    включаючи метадані відео та розбиття тексту на частини
    """
    file_path = str(file_path)  # Перетворення шляху на рядок
    original_uuid = hashlib.sha256(txt.encode()).hexdigest()  # Створення унікального ID
    doc_id = f"vid_{meta['video_id']}"  # Детермінований ID
    chunks = chunk_text(txt, size=chunk_size, overlap=overlap)  # Розбиття тексту на частини
    upload_date = dt.now().isoformat()  # Поточна дата та час

    return {
        "doc_id": doc_id,
        "original_uuid": original_uuid,
        "title": meta.get("title", ""),
        "video_meta": {                               # всі метадані відео
            "video_id":     meta["video_id"],
            "title":        meta.get("title", ""),
            
        },
        "content": txt,
        "chunks": [
            {
                "chunk_id":       f"{doc_id}_chunk_{i}",
                "original_index": i,
                "content":        c,
                "file_path":      file_path,
                "file_name":      original_filename,
                "upload_date":    upload_date,
            }
            for i, c in enumerate(chunks)
        ],
    }




def extract_audio_from_video(video_path: str) -> bytes:
    """
    Run FFmpeg to decode the audio track to 16 kHz, mono, 16-bit PCM,
    then convert to float32 in [-1.0, +1.0] for Whisper.
    """
    cmd = [
        "ffmpeg", "-nostdin", "-y",
        "-i", video_path,
        "-f", "s16le", "-ac", "1", "-acodec", "pcm_s16le", "-ar", "16000",
        "pipe:1"
    ]
    # Capture stdout, suppress stderr noise
    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.DEVNULL)
    pcm_bytes, _ = proc.communicate()
    if proc.returncode != 0:
        raise RuntimeError(f"FFmpeg exited with code {proc.returncode}")

    # Interpret bytes as int16 little-endian, then scale to float32
    audio_int16 = np.frombuffer(pcm_bytes, dtype=np.int16)
    audio_float = audio_int16.astype(np.float32) / 32768.0
    return audio_float

# Функції для екстракції даних з різних типів файлів
def data_extraction(path):
    """Визначає тип файлу за розширенням та викликає відповідну функцію екстракції"""
    ext = path.rsplit(".", 1)[-1].lower()  # Отримання розширення файлу
    if ext == "pdf":
        return extract_from_pdf(path)  # Екстракція з PDF
    if ext in ("xlsx", "xls"):
        return extract_from_excel(path)  # Екстракція з Excel
    if ext in ("doc", "docx"):
        return extract_from_word(path)  # Екстракція з Word
    if ext == "md":
        return extract_from_markdown(path)  # Екстракція з Markdown
    if ext == "csv":
        return extract_from_csv(path)  # Екстракція з CSV
    raise ValueError(f"Неподдерживаемый формат: {ext}")  # Виклик помилки для непідтримуваних форматів

def extract_from_pdf(p):
    """Витягує текст з PDF-файлу, об'єднуючи текст з усіх сторінок"""
    txt = ""  # Порожній рядок для зберігання тексту
    with open(p, "rb") as f:
        for pg in PdfReader(f).pages:  # Перебір сторінок
            txt += pg.extract_text() or ""  # Додавання тексту сторінки
    return txt  # Повернення тексту

def extract_from_excel(p, **kw):
    """Витягує дані з Excel-файлу, перетворюючи їх на DataFrame"""
    return pd.read_excel(p, dtype=str, **kw).fillna("")  # Читання Excel-файлу та заповнення пропусків

def extract_from_word(p):
    """Витягує текст з Word-документа, об'єднуючи текст з усіх параграфів"""
    return "\n".join(par.text for par in DocxDocument(p).paragraphs)  # Об'єднання тексту параграфів

def extract_from_markdown(p):
    """Витягує текст з Markdown-файлу, читаючи його як текстовий файл"""
    return open(p, encoding="utf-8").read()  # Читання файлу

def extract_from_csv(p, **kw):
    """Витягує дані з CSV-файлу, перетворюючи їх на DataFrame"""
    return pd.read_csv(p, dtype=str, **kw).fillna("")  # Читання CSV-файлу та заповнення пропусків

def prepare_text(text: Any) -> str:
    """
    Підготовка тексту (DataFrame, dict, або інший тип) до обробки, перетворюючи його на чистий рядок
    """
    # 1) Якщо DataFrame — склеюємо всі комірки
    if isinstance(text, pd.DataFrame):
        # об'єднуємо всі значення в один рядок
        text = ' '.join(text.values.flatten().astype(str))  # Об'єднання всіх значень
    # 2) Якщо dict — склеюємо всі значення
    elif isinstance(text, dict):
        text = ' '.join(str(v) for v in text.values())  # Об'єднання всіх значень
    # 3) Якщо не рядок — приводимо до рядка
    elif not isinstance(text, str):
        text = str(text)  # Перетворення на рядок

    # 4) Очищаємо отриманий рядок
    # видаляємо переведення рядків, табуляції та зайві пробіли
    txt = text.replace('\r', ' ').replace('\n', ' ').replace('\t', ' ')  # Заміна спеціальних символів
    txt = ' '.join(txt.split())  # прибираємо подвійні/зайві пробіли

    return txt.strip()  # Повернення очищеного рядка

def chunk_text(text: str, size=st.session_state.chunk_size, overlap=st.session_state.chunk_overlap):
    """Розбиває текст на частини вказаного розміру з перекриттям для кращого пошуку"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=size,  # Розмір частини
        chunk_overlap=overlap,  # Перекриття між частинами
        length_function=len,  # Функція для визначення довжини
        is_separator_regex=False  # Без використання регулярних виразів
    )
    return splitter.split_text(text)  # Розбиття тексту


def build_json(text: str, original_filename: str, path: str) -> dict:
    """Створює JSON-структуру для тексту з метаданими для зберігання в Milvus"""
    original_uuid = hashlib.sha256(text.encode()).hexdigest()  # Створення унікального ID
    doc_id = original_uuid[:16]  # Фіксований ідентифікатор (перші 16 символів)
    chunks = chunk_text(text)  # Розбиття тексту на частини
    upload_date = dt.now().isoformat()  # Поточна дата та час
    return {
        "doc_id": doc_id,  # ID документа
        "original_uuid": original_uuid,  # Повний унікальний ID
        "content": text,  # Повний текст
        "chunks": [
            {
                "chunk_id": f"{doc_id}_c{i}",  # ID частини (детермінований)
                "original_index": i,  # Індекс частини
                "content": c,  # Текст частини
                "file_path": path,  # Шлях до файлу
                "file_name": original_filename,  # Оригінальна назва файлу
                "upload_date": upload_date,  # Дата завантаження
            }
            for i, c in enumerate(chunks)  # Перебір частин
        ],
    }


def build_image_json(text: str, file_path: str, original_filename: str) -> dict:
    """Створює JSON-структуру для зображення з описом та метаданими для зберігання в Milvus"""
    txt = prepare_text(text)  # Підготовка тексту

    original_uuid = hashlib.sha256(txt.encode()).hexdigest()  # Створення унікального ID
    doc_id = original_uuid[:16]  # Фіксований ідентифікатор (перші 16 символів)
    chunks = chunk_text(txt)  # Розбиття тексту на частини
    upload_date = dt.now().isoformat()  # Поточна дата та час

    return {
        "doc_id": doc_id,  # ID документа
        "original_uuid": original_uuid,  # Повний унікальний ID
        "content": txt,  # Повний текст
        "chunks": [
            {
                "chunk_id": f"{doc_id}_c{i}",  # ID частини (детермінований)
                "content": c,  # Текст частини
                "file_path": file_path,  # Шлях до файлу
                "file_name": original_filename,  # Оригінальна назва файлу
                "upload_date": upload_date,  # Дата завантаження
                "original_index": i,  # Індекс частини
            }
            for i, c in enumerate(chunks)  # Перебір частин
        ],
    }


def init_db():
    """
    Ініціалізує базу даних, створюючи таблицю history, якщо вона не існує.
    Таблиця містить поля:
    - id: унікальний ідентифікатор запису
    - ts: часова мітка взаємодії
    - mode: режим взаємодії (документ, зображення, відео, чат)
    - query: запит користувача
    - answer: відповідь системи
    """
    conn = sqlite3.connect(DB_PATH)  # Підключення до бази даних
    cur = conn.cursor()  # Створення курсора для виконання SQL-запитів
    cur.execute(
        """CREATE TABLE IF NOT EXISTS history(
                id     INTEGER PRIMARY KEY AUTOINCREMENT,  
                ts     TEXT,  
                mode   TEXT,  
                query  TEXT,  
                answer TEXT   
        )"""
    )
    conn.commit(); conn.close()  # Збереження змін та закриття з'єднання

def get_conn():
    """Створює з'єднання з базою даних та ініціалізує таблицю, якщо потрібно"""
    conn = sqlite3.connect(DB_PATH)
    # створюємо таблицю один раз
    conn.execute(
        """CREATE TABLE IF NOT EXISTS history (
               id     INTEGER PRIMARY KEY AUTOINCREMENT,
               ts     TEXT,
               mode   TEXT,
               query  TEXT,
               answer TEXT
           )"""
    )
    return conn

def log_interaction(mode, query, answer):
    """Записує взаємодію користувача з системою в базу даних"""
    ts = datetime.now().isoformat(timespec="seconds")
    with get_conn() as conn:
        conn.execute(
            "INSERT INTO history(ts,mode,query,answer) VALUES (?,?,?,?)",
            (ts, mode, query, answer),
        )
@st.cache_resource(show_spinner="Завантажуємо модель Whisper…")
def get_whisper(model_size: str = "base"):
    return whisper.load_model(model_size, device="cuda" if torch.cuda.is_available() else "cpu")

def build_audio_json(
    txt: str,
    meta: Dict[str, Any],
    file_path: str,
    original_filename: str,
    chunk_size: int = st.session_state.chunk_size,
    overlap: int = st.session_state.chunk_overlap,
) -> Dict[str, Any]:
    """
    Формирует единый JSON для аудио:
    - разбивает prepared-text на чанки
    - собирает метаданные
    - возвращает структуру, готовую к upsert в Milvus
    """
    file_path = str(file_path)
    original_uuid = hashlib.sha256(txt.encode()).hexdigest()
    doc_id = f"aud_{meta['unique_audio_id']}"
    chunks = chunk_text(txt, size=chunk_size, overlap=overlap)

    return {
        "doc_id": doc_id,
        "original_uuid": original_uuid,
        "audio_meta": {
            "audio_id":         meta["audio_id"],
            "unique_audio_id":  meta["unique_audio_id"],
            "title":            meta["title"],
            "upload_date":      meta["upload_date"],
        },
        "content": txt,
        "chunks": [
            {
                "chunk_id":       f"{doc_id}_chunk_{i}",
                "original_index": i,
                "content":        c,
                "file_path":      file_path,
                "file_name":      original_filename,
                "upload_date":    meta["upload_date"],
            }
            for i, c in enumerate(chunks)
        ],
    }


def process_audio(filepath: str, model_size: str, dir_name: str, title: str) -> str:
    """
    Транскрибує аудіо і повертає чистый текст.
    Ничего больше не делает — не разбивает на чанки и не формирует метаданные.
    """
    # Загрузили модель один раз (кешируется get_whisper)
    whisper_model = get_whisper(model_size)
    logger = logging.getLogger(__name__)
    logger.info(f"Транскрибація аудіо: {filepath}")

# --- Транскрибація аудіо ---
#Для того аби змусити модель повертати весь текст, навіть якщо там на початку є музіка, потрібно щоб з початку файлу був голос а не музика та щоб формат файлу ьув якісний (wav):
    result = whisper_model.transcribe(str(filepath))
    placeholder = st.empty()
    transcript = result["text"]
    placeholder.text(transcript) 
    # Сохраняем транскрипт рядом с файлом
    transcript_path = os.path.join(dir_name, f"{title}_transcript.txt")
    with open(transcript_path, "w", encoding="utf-8") as f:
        f.write(transcript)
    logger.info(f"Транскрипт збережено: {transcript_path}")
    return transcript


def audio_mode(collection_name: str, summary: bool):
    audio_dir = "audio"
    create_dir(audio_dir)

    uploaded_file = st.file_uploader(
        "Оберіть аудіофайл",
        type=["mp3", "wav", "flac", "aac", "m4a", "ogg"],
    )
    if uploaded_file is None:
        return

    # --- Сохранение файла ---
    full_name = uploaded_file.name
    title, ext = os.path.splitext(full_name)
    filepath = os.path.join(audio_dir, full_name)
    with open(filepath, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Метаданные
    audio_id         = Path(filepath).stem
    safe_title       = clean_filename(title)
    unique_audio_id = f"{audio_id}_{hashlib.md5(filepath.encode()).hexdigest()[:8]}"
    upload_date      = dt.now().isoformat()

    meta = {
        "title":             title,
        "audio_id":          audio_id,
        "unique_audio_id":   unique_audio_id,
        "upload_date":       upload_date,
    }

    # Выбор модели и кнопка запуска
    model_size = st.selectbox(
        "Модель Whisper:",
        ["tiny", "base", "small", "medium", "large"],
        index=1,
    )
    if not st.button("Обробити"):
        return

    # --- Транскрибация ---
    transcript = process_audio(filepath, model_size, audio_dir, title)

    # Подготовка текста к чанкингу
    cleaned = prepare_text(transcript)
    st.session_state.audio_context_text['context'] = ""
    if summary and not st.session_state.audio_context_text.get("context"):
        llm = create_llm("qwen3:14b")
        summary = summarise_transcript(cleaned, llm)
        st.session_state.audio_context_text["context"] = summary

    # --- Построение единого JSON ---
    audio_json = build_audio_json(
        txt=cleaned,
        meta=meta,
        file_path=filepath,
        original_filename=full_name,
    )
    st.session_state["last_audio_json"] = audio_json
    st.session_state.audio_processed = True

    # --- Загрузка в Milvus ---
    dense_ef = create_bge_m3_embeddings()
    retriever = AudioHybridRetriever(
        client=st.session_state.milvus_client,
        collection_name=collection_name,
        dense_embedding_function=dense_ef,
    )
    if st.session_state.audio_processed:
        retriever.build_collection()
        for chunk in audio_json["chunks"]:
            metadata = {
                **audio_json["audio_meta"],
                "doc_id":        audio_json["doc_id"],
                "original_uuid": audio_json["original_uuid"],
                "chunk_id":      chunk["chunk_id"],
                "original_index": chunk["original_index"],
                "file_path":      chunk["file_path"],
                "file_name":      chunk["file_name"],
                "upload_date":    chunk["upload_date"],
            }
            retriever.upsert_data(chunk["content"], metadata)

    st.success("Аудіо оброблено та завантажено в колекцію.")
    st.session_state.audio_processed = True


def chat_audio_mode(collection_name, llm_option):
    if st.session_state.audio_context_text['context']:
        st.markdown(st.session_state.audio_context_text['context'])
    if not st.session_state.messages:
        st.session_state.messages.append({"role": "assistant", "content": "Привіт! Я готовий до розмови. Що вас цікавить?"})
         
    for m in st.session_state.messages:
        with st.chat_message(m["role"]): st.markdown(m["content"])
    placeholder = st.empty()

    # ввод користувача
    if q := st.chat_input("Ваш запит"):
        st.session_state.messages.append(dict(role="user", content=q))
        with st.chat_message("user"): st.markdown(q)
        with st.chat_message("assistant"):
            st.write("Шукаю відповідь…")
            dense_ef = create_bge_m3_embeddings()
            retr = AudioHybridRetriever(
            client=st.session_state.milvus_client,
            collection_name=collection_name,
            dense_embedding_function=dense_ef,) 
            retr.build_collection()
            chunks = retr.search(q, mode="hybrid", k=st.session_state.ret_k_results)
            ctx = "\n---\n".join([c["content"] for c in chunks]) if chunks else ""
            user_query = q+"\n---\n"+ctx
            if llm_option == "Україномовну":
                    # Формуємо єдиний рядок для генерації
                    system = UKR_SYSTEM_PROMPT
                    prompt_text = system + "\n\n" + user_query
                    # Викликаємо HuggingFace pipeline, який очікує рядок
                    gen = st.session_state.ukr_generator(
                            [
                                {"role": "user", "content": user_query}
                            ],
                            max_new_tokens=MAX_TOKENS,
                            do_sample=False)
                    # Результат — список з одним словником із ключем "generated_text"
                    answer = gen[0]["generated_text"]
            else:   
                placeholder = st.empty()
                answer = ""
                prompt = create_stream_prompt(st.session_state.system_prompt, q, ctx)
                model = st.session_state.llm_option
                for part in query_ollama(prompt, model):
                    chunk = part["message"]["content"]
                    # по‑символьно додаємо та одразу оновлюємо плейсхолдер
                    for ch in chunk:
                        answer += ch
                        placeholder.text(answer)   # або .markdown(caption)

            with st.expander("Показати використаний контекст"):
                for i in chunks:
                        st.markdown(
                            f"**Файл:** {i['file_name']}  \n"
                            f"**Chunk ID:** `{i['chunk_id']}`  \n"
                            f"**Дата завантаження:** {i['upload_date']}  \n"
                            f"**Score:** {i['score']:.4f}  \n\n"
                            f"> {i['content']}",
                        )
            log_interaction(st.session_state.current_mode, q, answer)
            st.session_state.messages.append(dict(role="assistant", content=answer))

def video_mode(collection_name: str, summary: bool):
    st.session_state.current_mode = "video"
    st.subheader("Обробка відео")
    col1, col2 = st.columns(2)
    with col1:
        url = st.text_input("Введіть посилання на відео в YouTube", key="video_url")
    with col2:
        uploaded_video = st.file_uploader(
            "Upload local video", type=["mp4","mov","avi","mpeg"],key="video_file")
    st.session_state.video_name = st.text_input('Введіть назву відео')
    model_size = st.selectbox("Whisper model",["tiny","base","small","medium","large"],index=1)

    if st.button("Обробка"):
        if uploaded_video is not None:
            video_dir = "video"
            Path(video_dir).mkdir(exist_ok=True)
            title = uploaded_video.name
            unique_video_id = f"{title}_{hashlib.md5(url.encode()).hexdigest()[:8]}"
            file_path = os.path.join(video_dir, uploaded_video.name)
            with open(file_path,"wb") as f: 
                f.write(uploaded_video.getbuffer())
            raw = extract_audio_from_video(file_path)            # NumPy array
            wav_path = os.path.join(video_dir, f"{title}.wav")
            sf.write(wav_path, raw, samplerate=16000)
            txt = process_audio(wav_path, model_size, video_dir, title)

        elif url:
            txt, file_path, title, unique_video_id = process_video(url,"video",model_size)
        txt_clean = prepare_text(txt)
        st.session_state.video_context_text['context'] = ""
        if summary and not st.session_state.video_context_text['context']:
            summary = summarise_transcript(txt_clean, create_llm(st.session_state.llm_option))
            st.session_state.video_context_text['context']=summary
        video_meta={"video_id": unique_video_id,
                    "title":title,
                    "url":url,
                    "file_path":file_path, 
                    "upload_date":dt.now().isoformat(),
                    }
        video_json=build_video_json(txt_clean,video_meta,file_path, title)
        dense_ef=create_bge_m3_embeddings()
        retr=VideoHybridRetriever(client=st.session_state.milvus_client, collection_name=collection_name,
                                   dense_embedding_function=dense_ef)
        retr.build_collection()
        for chunk in video_json["chunks"]:
            metadata = {
                **video_json["video_meta"],
                "doc_id":        video_json["doc_id"],
                "original_uuid": video_json["original_uuid"],
                "chunk_id":      chunk["chunk_id"],
                "original_index": chunk["original_index"],
                "file_path":      chunk["file_path"],
                "file_name":      chunk["file_name"],
                "upload_date":    chunk["upload_date"],
            }
            retr.upsert_data(chunk["content"], metadata)
        st.session_state.last_video_json=video_json
        st.session_state.video_processed=True
        st.success("Video processed and uploaded 🔄")

def chat_video_mode(collection_name, llm_option):
    if st.session_state.video_context_text['context']:
        st.markdown(st.session_state.video_context_text['context'])
    if not st.session_state.messages:
        st.session_state.messages.append({"role": "assistant", "content": "Привіт! Я готовий до розмови. Що вас цікавить?"})
         
    for m in st.session_state.messages:
        with st.chat_message(m["role"]): st.markdown(m["content"])
    placeholder = st.empty()

    # ввод користувача
 
    if q := st.chat_input("Ваш запит"):
     
        st.session_state.messages.append(dict(role="user", content=q))
        with st.chat_message("user"): st.markdown(q)
        with st.chat_message("assistant"):
            st.write("Шукаю відповідь…")
            dense_ef = create_bge_m3_embeddings()
            retr = VideoHybridRetriever(
            client=st.session_state.milvus_client,
            collection_name=collection_name,
            dense_embedding_function=dense_ef,) 
            retr.build_collection()
            chunks = retr.search(q, mode="hybrid", k=st.session_state.ret_k_results)
            ctx = "\n---\n".join([c["content"] for c in chunks]) if chunks else ""
            user_query = q+"\n---\n"+ctx
            #if llm_option == "Україномовну":
            #    chat = [
            #        {"role": "system", "content": "You are useful assistant. Use the info to answer user query. Answer in Ukrainian"},
            #        {"role": "user", "content": user_query}
            #        ]
            #    response = st.session_state.ukr_generator(chat, max_new_tokens=512)
            #    answer = response[0]["generated_text"][-1]["content"]
            if llm_option == "Україномовну":
                    # Формуємо єдиний рядок для генерації
                    system = UKR_SYSTEM_PROMPT
                    prompt_text = system + "\n\n" + user_query
                    # Викликаємо HuggingFace pipeline, який очікує рядок
                    gen = st.session_state.ukr_generator(
                            [
                                {"role": "user", "content": user_query}
                            ],
                            max_new_tokens=512,
                            do_sample=False)
                    # Результат — список з одним словником із ключем "generated_text"
                    answer = gen[0]["generated_text"]
            else:   
                placeholder = st.empty()
                answer = ""
                prompt = create_stream_prompt(st.session_state.system_prompt, q, ctx)
                model = st.session_state.llm_option
                for part in query_ollama(prompt, model):
                    chunk = part["message"]["content"]
                    # по‑символьно додаємо та одразу оновлюємо плейсхолдер
                    for ch in chunk:
                        answer += ch
                        placeholder.text(answer)   # або .markdown(caption)
            with st.expander("Показати використаний контекст"):
                for i in chunks:
                        st.markdown(
                            f"**Файл:** {i['file_name']}  \n"
                            f"**Chunk ID:** `{i['chunk_id']}`  \n"
                            f"**Дата завантаження:** {i['upload_date']}  \n"
                            f"**Score:** {i['score']:.4f}  \n\n"
                            f"> {i['content']}",
                        )
            log_interaction(st.session_state.current_mode, q, answer)
            st.session_state.messages.append(dict(role="assistant", content=answer))

def image_mode(collection_name, summary = True):
    st.session_state.current_mode = "image"
    st.subheader("Завантаження зображень")
    uploaded_images = st.file_uploader(
    "Обери одне або кілька зображень",
    type=["png", "jpg", "jpeg", "webp"],
    accept_multiple_files=True,
    key="image_uploader"
)
    if uploaded_images:
        image_dir = "images"
        if not os.path.exists(image_dir):
            os.makedirs(image_dir)
            st.info(f"Створено директорію {image_dir} для збереження зображень")
        
        for img_file in uploaded_images:
            # превью в интерфейсе
            st.image(img_file)
            pil_img = Image.open(img_file)
            img_b64 = [image_to_base64(pil_img)]
            prompt = IMAGE_DESCRIPTION_SYSTEM_PROMPT
            with st.spinner("Генерую опис…"):
                placeholder = st.empty()
                caption = ""

                # ollama.chat(..., stream=True) повертає генератор,
                #   у part["message"]["content"] — черговий фрагмент тексту
                for part in query_ollama(prompt, IMAGE_DESCRIPTION_MODEL, img_b64):
                    chunk = part["message"]["content"]
                    # по‑символьно додаємо та одразу оновлюємо плейсхолдер
                    for ch in chunk:
                        caption += ch
                        placeholder.text(caption)   # або .markdown(caption)
                llm = create_llm(st.session_state.llm_option)
                print(llm, 'llm')
                st.session_state.image_context_text['context'] = ""
                if summary and not st.session_state.image_context_text['context']:
                    summary  = summarise_transcript(caption, llm)
                    summary = remove_think(summary) # видаляє дужки з тексту
                    st.session_state.image_context_text['context'] = summary
  
            # ---------- сохранение ----------
            file_ext = os.path.splitext(img_file.name)[1]
            print(file_ext, 'file_ext')
            unique_name = f"{uuid.uuid4().hex}{file_ext}"

            print(unique_name, 'unique_name')
                    
            # Зберігаємо копію в директорії images
            image_path = os.path.join(image_dir, unique_name)
            with open(image_path, "wb") as f:
                f.write(img_file.getbuffer())
            st.success(f"Копію файлу збережено в {image_path}")
            st.session_state.image_processed = True

            # ---------- запись в Milvus ----------
            raw = build_image_json(caption, image_path, img_file.name)
            if isinstance(raw, dict):
                docs = [raw]
            elif isinstance(raw, list):
                docs = raw
            else:
                raise ValueError("Неподдерживаемый формат из build_json")
            
            print(docs, 'docs')
            
            is_insert = True
            if is_insert:
                dense_ef = create_bge_m3_embeddings()
                standard_retriever = ImageHybridRetriever(
                    client=st.session_state.milvus_client,
                    collection_name=collection_name,
                    dense_embedding_function=dense_ef,)
                standard_retriever.build_collection()          # ← додали
                for doc in docs:
                #Added funct
                    existing = standard_retriever.client.query(
                    collection_name=collection_name,
                    filter=f'original_uuid == "{doc["original_uuid"]}"',
                    output_fields=["chunk_id"])
                    existing_ids = {row["chunk_id"] for row in existing}
                #End of adding

                    for chunk in doc["chunks"]:
                        #Added funct
                        if chunk["chunk_id"] in existing_ids:        # дубль – пропускаем
                            continue
                        #End of adding
                        metadata = {
                            "doc_id":          doc["doc_id"],
                            "original_uuid":  doc["original_uuid"],
                            "chunk_id":       chunk["chunk_id"],
                            "original_index": chunk["original_index"],
                            "content":        chunk["content"],
                            "file_path":      chunk["file_path"],
                            "file_name":      chunk["file_name"],
                            "upload_date":    chunk["upload_date"],
                        }
                        standard_retriever.upsert_data(chunk["content"], metadata)
            st.session_state.image_processed=True

def chat_image_mode(collection_name, llm_option):
    if st.session_state.image_context_text['context']:
        st.markdown(st.session_state.image_context_text['context'])
    if not st.session_state.messages:
        st.session_state.messages.append({"role": "assistant", "content": "Привіт! Я готовий до розмови. Що вас цікавить?"})
         
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
    placeholder = st.empty()

    # Ввод пользователя
    if q := st.chat_input("Ваш запит"):
        st.session_state.messages.append(dict(role="user", content=q))
        with st.chat_message("user"): 
            st.markdown(q)
        with st.chat_message("assistant"):
            st.write("Шукаю відповідь…")
            dense_ef = create_bge_m3_embeddings()

            retriever = ImageHybridRetriever(
                client=st.session_state.milvus_client,
                collection_name=collection_name,
                dense_embedding_function=dense_ef,
            )
            retriever.build_collection()
            results = retriever.search(q, mode="hybrid", k=st.session_state.ret_k_results)
            if not results:
                answer = "На жаль, не знайшов релевантної інформації."
            else:
                ctx = "\n---\n".join([c["content"] for c in results]) if results else ""
                user_query = q+"\n---\n"+ctx
                if llm_option == "Україномовну":
                        # Формуємо єдиний рядок для генерації
                        system = UKR_SYSTEM_PROMPT
                        prompt_text = system + "\n\n" + user_query
                        # Викликаємо HuggingFace pipeline, який очікує рядок
                        gen = st.session_state.ukr_generator(
                                [
                                    {"role": "user", "content": user_query}
                                ],
                                max_new_tokens=MAX_TOKENS,
                                do_sample=False)
                        # Результат — список з одним словником із ключем "generated_text"
                        answer = gen[0]["generated_text"]
                else:
                    placeholder = st.empty()
                    answer = ""
                    prompt = create_stream_prompt(st.session_state.system_prompt, q, ctx)
                    model = st.session_state.llm_option
                    for part in query_ollama(prompt, model):
                        chunk = part["message"]["content"]
                        # по‑символьно додаємо та одразу оновлюємо плейсхолдер
                        for ch in chunk:
                            answer += ch
                            placeholder.text(answer)   # або .markdown(caption)
                with st.expander("Показати використаний контекст"):
                    for i in results:
                            st.markdown(
                                f"**Файл:** {i['file_name']}  \n"
                                f"**Chunk ID:** `{i['chunk_id']}`  \n"
                                f"**Дата завантаження:** {i['upload_date']}  \n"
                                f"**Score:** {i['score']:.4f}  \n\n"
                                f"> {i['content']}",
                            )
                log_interaction(st.session_state.current_mode, q, answer)
                st.session_state.messages.append(dict(role="assistant", content=answer))
                

def main_chat(collection_name):
    if "messages" not in st.session_state: 
        st.session_state.messages = []  # Ініціалізація повідомлень
    # Якщо чат тільки що запущений, відправляємо привітання
    if not st.session_state.messages:
        st.session_state.messages.append({
            "role": "assistant",
            "content": "Привіт! Я ваш чат‑агент. Я вмію шукати по збережених векторах у Milvus і давати відповіді. Що запитуєте?"
        })

    # Відображення історії повідомлень
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):  # Створення блоку повідомлення
            st.markdown(msg["content"])  # Відображення тексту повідомлення
    placeholder = st.empty()
    # Введення користувача
    if query := st.chat_input("Введіть ваше запитання...", key="chat_input"):
        # Додавання повідомлення користувача в історію
        st.session_state.messages.append({"role": "user", "content": query})
        with st.chat_message("user"):  # Створення блоку повідомлення користувача
            st.markdown(query)  # Відображення запиту користувача

        # Запуск пошуку по векторах
        with st.chat_message("assistant"):  # Створення блоку повідомлення асистента
            st.write("Шукаю відповідь...")  # Відображення статусу пошуку
            # Створення функції вбудовування
            dense_ef = create_bge_m3_embeddings()
            retriever = HybridRetriever(  # Створення гібридного пошуковика
                client=st.session_state.milvus_client,  # Клієнт Milvus
                collection_name=collection_name,  # Назва колекції
                dense_embedding_function=dense_ef,  # Функція вбудовування
            )
            retriever.build_collection()
            stats = st.session_state.milvus_client.get_collection_stats(st.session_state.collection_name)
            entity_count = int(stats["row_count"])
            if entity_count == 0:
                answer = "На жаль, не знайшов релевантної інформації."
                metadata = {
                        "doc_id":          'none',
                        "original_uuid":  'none',
                        "chunk_id":      'none',
                        "original_index": 0,
                        "content":        'none',
                        "file_path":      'none',
                        "file_name":      'none',
                        "upload_date":    'none',
                    }

                retriever.upsert_data(answer, metadata)
            results = retriever.search(query, mode="hybrid", k=st.session_state.ret_k_results)
            if not results:  # Якщо результати не знайдені
                answer = "На жаль, не знайшов релевантної інформації."
            else:
                best = "\n".join(r["content"] for r in results) if results else ""
                user_query = query+"\n---\n"+best  # Формування запиту з контекстом
                if st.session_state.llm_option == "Україномовну":  # Якщо вибрана україномовна модель
                    chat = [
                        {"role": "system", "content": "You are useful assistant. Use the info to answer user query. Answer in Ukrainian"},
                        {"role": "user", "content": user_query}
                        ]
                    response = st.session_state.ukr_generator(chat, max_new_tokens=512)  # Генерація відповіді
                    answer = response[0]["generated_text"][-1]["content"]  # Отримання тексту відповіді
                else:
                # Формування промпту та виклик LLM
                    placeholder = st.empty()
                    answer = ""
                    prompt = create_stream_prompt(st.session_state.system_prompt, query, best)
                    model = st.session_state.llm_option
                    for part in query_ollama(prompt, model):
                        chunk = part["message"]["content"]
                        # по‑символьно додаємо та одразу оновлюємо плейсхолдер
                        for ch in chunk:
                            answer += ch
                            placeholder.text(answer) 
            with st.expander("Показати використаний контекст"):  # Створення розгортаємого блоку
                    for i in results:
                        st.markdown(
                            f"**Файл:** {i['file_name']}  \n"  # Відображення імені файлу
                            f"**Chunk ID:** `{i['chunk_id']}`  \n"  # Відображення ID фрагмента
                            f"**Дата завантаження:** {i['upload_date']}  \n"  # Відображення дати завантаження
                            f"**Score:** {i['score']:.4f}  \n\n"  # Відображення оцінки релевантності
                            f"> {i['content']}"  # Відображення вмісту фрагмента
                            )
            
            st.session_state.messages.append({"role": "assistant", "content": answer})  # Додавання відповіді в історію
            log_interaction(st.session_state.current_mode or "chat", query, answer)  # Логування взаємодії

def document_mode(collection_name, summary):
    uploaded_file = st.file_uploader("Завантажте документ", type=['pdf', 'xlsx', 'xls', 'doc', 'docx', 'md'])
    if uploaded_file is not None:
        ext = os.path.splitext(uploaded_file.name)[1].lower()
        if ext in [".xlsx", ".xls", ".csv"]:
            if ext == ".csv":
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            st.dataframe(df.head(10), use_container_width=True)
            preview_text = " ".join(df.astype(str).head(10).values.flatten())
        elif ext == ".pdf":
            reader = PdfReader(uploaded_file)
            first_page_text = reader.pages[0].extract_text() or ""
            preview_text = first_page_text[:1000]
        elif ext in [".docx", ".doc"]:
            docx_obj = DocxDocument(uploaded_file)
            text = "\n".join(p.text for p in docx_obj.paragraphs)
            preview_text = text[:1000]
        else:  # docx, md, txt, etc.
            raw_bytes = uploaded_file.read()
            text = raw_bytes.decode("utf-8", errors="ignore")
            preview_text = text[:1000]
        uploaded_file.seek(0)
        # Отримуємо розширення файлу з оригінальної назви
        file_extension = os.path.splitext(uploaded_file.name)[1]
        # Генеруємо унікальне ім'я файлу використовуючи UUID
        unique_filename = f"{uuid.uuid4().hex}{file_extension}"
        # Формуємо повний шлях до файлу в папці data
        file_path = os.path.join("data", unique_filename)
        # Зберігаємо оригінальну назву файлу
        original_filename = uploaded_file.name
        
        # Відкриваємо файл для запису в бінарному режимі
        with open(file_path, "wb") as f:
            # Записуємо вміст завантаженого файлу
            f.write(uploaded_file.getbuffer())
        # Показуємо повідомлення про успішне збереження
        st.success(f"Файл {original_filename} успішно збережено як {file_path}")
        raw = data_extraction(file_path)
        raw_str = raw.to_csv(index=False) if isinstance(raw, pd.DataFrame) else (
        json.dumps(raw, ensure_ascii=False) if isinstance(raw, dict) else str(raw)
    )
        txt = prepare_text(raw_str)
        st.markdown(f"**Перші 1000 символів файлу:** \n\n {txt[:1000]}")
        st.session_state.document_context_text['context'] = ""
        if summary and not st.session_state.document_context_text['context']:
            summary = summarise_transcript(txt, create_llm(st.session_state.llm_option))
            summary = remove_think(summary) # видаляє дужки з тексту
            st.session_state.document_context_text['context'] = summary
        raw = build_json(txt, original_filename, file_path)

        if isinstance(raw, dict):
            docs = [raw]
        elif isinstance(raw, list):
            docs = raw
        else:
            raise ValueError("Неподдерживаемый формат из build_json")

        dense_ef = create_bge_m3_embeddings()
        standard_retriever = HybridRetriever(
        client=st.session_state.milvus_client,
        collection_name=collection_name,
        dense_embedding_function=dense_ef,
        )

        is_insert = True
        if is_insert:
            standard_retriever.build_collection()
            for doc in docs:
                existing = standard_retriever.client.query(
                collection_name=collection_name,
                filter=f'original_uuid == "{doc["original_uuid"]}"',
                output_fields=["chunk_id"])
                existing_ids = {row["chunk_id"] for row in existing}
                for chunk in doc["chunks"]:
                    if chunk["chunk_id"] in existing_ids:        # дубль – пропускаем
                        continue
                    metadata = {
                        "doc_id":          doc["doc_id"],
                        "original_uuid":  doc["original_uuid"],
                        "chunk_id":       chunk["chunk_id"],
                        "original_index": chunk["original_index"],
                        "content":        chunk["content"],
                        "file_path":      chunk["file_path"],
                        "file_name":      chunk["file_name"],
                        "upload_date":    chunk["upload_date"],
                    }
                    standard_retriever.upsert_data(chunk["content"], metadata)
        st.session_state.document_processed = True
        st.session_state.document_start = False

        
def chat_document_mode(collection_name, llm_option):
    if st.session_state.document_context_text['context']:
        st.markdown(st.session_state.document_context_text['context'])
    st.session_state.doc_mode = "chat"
    if not st.session_state.messages:
        st.session_state.messages.append({"role": "assistant", "content": "Привіт! Я готовий до розмови. Що вас цікавить?"})
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
    placeholder = st.empty()
    # Ввод пользователя
    if query := st.chat_input("Введіть ваше запитання...", key="document_chat_input"):
    # Добавляем сообщение пользователя в историю
        st.session_state.messages.append({"role": "user", "content": query})
        with st.chat_message("user"):
            st.markdown(query)

        # Стартуем поиск по векторам
        with st.chat_message("assistant"):
            st.write("Шукаю відповідь...")
            # Создаём извлекатель векторов
            dense_ef = create_bge_m3_embeddings()
            retriever = HybridRetriever(
                client=st.session_state.milvus_client,
                collection_name=collection_name,
                dense_embedding_function=dense_ef,
            )
            retriever.build_collection()          # ← додали
            # Ищем 5 наиболее похожих чанков
            results = retriever.search(query, mode="hybrid", k=st.session_state.ret_k_results)
            if not results:
                answer = "На жаль, не знайшов релевантної інформації."
            else:
                best = ""
                for i in results:
                    best += i["content"]
                # Формируем промпт и вызываем LLM
                prompt = create_prompt(st.session_state.system_prompt)
                if llm_option == "Україномовну":
                    # Формуємо єдиний рядок для генерації
                    system = UKR_SYSTEM_PROMPT
                    prompt_text = system + "\n\n" + query + "\n---\n" + best
                    # Викликаємо HuggingFace pipeline, який очікує рядок
                    gen = st.session_state.ukr_generator(prompt_text, max_new_tokens=512, do_sample=False)
                    # Результат — список з одним словником із ключем "generated_text"
                    answer = gen[0]["generated_text"]
                    st.markdown(f"**Answer:** {answer}")
                    st.markdown(f"**lENG:** {len(answer)}")
                else:   
                    placeholder = st.empty()
                    answer = ""
                    prompt = create_stream_prompt(st.session_state.system_prompt, query, best)
                    model = st.session_state.llm_option
                    for part in query_ollama(prompt, model):
                        chunk = part["message"]["content"]
                        chunk = remove_think(chunk)
                        for ch in chunk:
                            answer += ch
                            placeholder.text(answer)   # або .markdown(caption)
            
            with st.expander("Показати використаний контекст"):
                        for i in results:
                            st.markdown(
                                f"**Файл:** {i['file_name']}  \n"
                                f"**Chunk ID:** `{i['chunk_id']}`  \n"
                                f"**Дата завантаження:** {results[0]['upload_date']}  \n"
                                f"**Score:** {i['score']:.4f}  \n\n"
                                f"> {i['content']}"
                            )
            st.session_state.messages.append({"role": "assistant", "content": answer})
            log_interaction(st.session_state.current_mode or "chat", query, answer)

def create_summary_prompt() -> ChatPromptTemplate:
    return ChatPromptTemplate.from_messages([
        ("system", SUMMARY_SYSTEM_PROMPT),
        ("human", "{text}")
    ])

def summarise_transcript(transcript: str, llm) -> str:
    prompt = create_summary_prompt()
    chain  = create_chain(llm, prompt)    # create_chain = prompt | llm
    response = chain.invoke({"text": transcript})
    return response.content

def on_summarise_video():
    st.session_state.video_summary = True

def on_summarise_audio():
    """Викликається при натисканні кнопки — робить summary і зберігає в сесії."""
    st.session_state.audio_summary = True

def on_summarise_image():
    """Викликається при натисканні кнопки — робить summary і зберігає в сесії."""
    st.session_state.image_summary = True

def on_summarise_document():
    """Викликається при натисканні кнопки — робить summary і зберігає в сесії."""
    st.session_state.document_summary = True
    
